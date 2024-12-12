import tkinter as tk
from tkinter import filedialog, colorchooser, font, messagebox
from tkinter.scrolledtext import ScrolledText
from docx import Document
from fpdf import FPDF
import os

class TextEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Text Editor")
        self.file_path = None

        # Default theme settings
        self.theme = {
            "bg_color": "white",
            "fg_color": "black",
            "font": ("Arial", 12)
        }

        # Text area
        self.text_area = ScrolledText(root, wrap=tk.WORD, font=self.theme["font"], bg=self.theme["bg_color"], fg=self.theme["fg_color"])
        self.text_area.pack(expand=1, fill=tk.BOTH)

        # Menu bar
        self.menu = tk.Menu(root)
        self.root.config(menu=self.menu)

        # File menu
        file_menu = tk.Menu(self.menu, tearoff=0)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save as PDF", command=self.save_as_pdf)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=root.quit)
        self.menu.add_cascade(label="File", menu=file_menu)

        # Format menu
        format_menu = tk.Menu(self.menu, tearoff=0)
        format_menu.add_command(label="Font", command=self.change_font)
        format_menu.add_command(label="Font Color", command=self.change_font_color)
        format_menu.add_command(label="Background Color", command=self.change_bg_color)
        self.menu.add_cascade(label="Format", menu=format_menu)

        # Customization menu
        customize_menu = tk.Menu(self.menu, tearoff=0)
        customize_menu.add_command(label="Set Theme", command=self.set_theme)
        self.menu.add_cascade(label="Customize", menu=customize_menu)

    def new_file(self):
        self.text_area.delete(1.0, tk.END)
        self.file_path = None

    def open_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Documents", "*.docx"), ("Text Files", "*.txt")]
        )
        if file_path:
            self.file_path = file_path
            if file_path.endswith(".docx"):
                doc = Document(file_path)
                self.text_area.delete(1.0, tk.END)
                for para in doc.paragraphs:
                    self.text_area.insert(tk.END, para.text + "\n")
            elif file_path.endswith(".txt"):
                with open(file_path, "r") as file:
                    content = file.read()
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, content)

    def save_file(self):
        if not self.file_path:
            self.file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("Text Files", "*.txt")],
            )
        if self.file_path:
            if self.file_path.endswith(".docx"):
                doc = Document()
                for line in self.text_area.get(1.0, tk.END).splitlines():
                    doc.add_paragraph(line)
                doc.save(self.file_path)
            elif self.file_path.endswith(".txt"):
                with open(self.file_path, "w") as file:
                    file.write(self.text_area.get(1.0, tk.END).strip())

    def save_as_pdf(self):
        pdf_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if pdf_path:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in self.text_area.get(1.0, tk.END).splitlines():
                pdf.cell(0, 10, txt=line, ln=True)
            pdf.output(pdf_path)

    def change_font(self):
        font_choice = font.Font(font=self.text_area['font'])
        font_dialog = tk.Toplevel(self.root)
        font_dialog.title("Choose Font")

        font_family_label = tk.Label(font_dialog, text="Font Family:")
        font_family_label.pack()

        font_family_var = tk.StringVar(value=font_choice.actual()["family"])
        font_family_entry = tk.Entry(font_dialog, textvariable=font_family_var)
        font_family_entry.pack()

        font_size_label = tk.Label(font_dialog, text="Font Size:")
        font_size_label.pack()

        font_size_var = tk.IntVar(value=font_choice.actual()["size"])
        font_size_spinbox = tk.Spinbox(font_dialog, from_=8, to=72, textvariable=font_size_var)
        font_size_spinbox.pack()

        def apply_font():
            new_font = (font_family_var.get(), font_size_var.get())
            self.text_area.config(font=new_font)
            self.theme["font"] = new_font
            font_dialog.destroy()

        apply_button = tk.Button(font_dialog, text="Apply", command=apply_font)
        apply_button.pack()

    def change_font_color(self):
        color_code = colorchooser.askcolor(title="Choose font color")[1]
        if color_code:
            self.text_area.config(fg=color_code)
            self.theme["fg_color"] = color_code

    def change_bg_color(self):
        color_code = colorchooser.askcolor(title="Choose background color")[1]
        if color_code:
            self.text_area.config(bg=color_code)
            self.theme["bg_color"] = color_code

    def set_theme(self):
        theme_dialog = tk.Toplevel(self.root)
        theme_dialog.title("Set Theme")

        bg_label = tk.Label(theme_dialog, text="Background Color:")
        bg_label.pack()

        bg_button = tk.Button(theme_dialog, text="Choose", command=lambda: self.update_theme_color("bg_color"))
        bg_button.pack()

        fg_label = tk.Label(theme_dialog, text="Font Color:")
        fg_label.pack()

        fg_button = tk.Button(theme_dialog, text="Choose", command=lambda: self.update_theme_color("fg_color"))
        fg_button.pack()

        apply_button = tk.Button(theme_dialog, text="Apply", command=theme_dialog.destroy)
        apply_button.pack()

    def update_theme_color(self, color_type):
        color_code = colorchooser.askcolor(title="Choose Color")[1]
        if color_code:
            if color_type == "bg_color":
                self.text_area.config(bg=color_code)
                self.theme["bg_color"] = color_code
            elif color_type == "fg_color":
                self.text_area.config(fg=color_code)
                self.theme["fg_color"] = color_code

if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditorApp(root)
    root.mainloop()
